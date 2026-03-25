"""
本地 QA 冒烟：不依赖 pytest，从项目根执行：
  python tests/qa_smoke.py
覆盖：首页、诊断、docx 上传（含正负例）、论文.txt 写入校验。
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from fastapi.testclient import TestClient

from server import THESIS_PATH, app


def make_minimal_docx_bytes() -> bytes:
    from docx import Document

    buf = io.BytesIO()
    d = Document()
    d.add_paragraph("QA_MARKER_PARA_一")
    d.add_paragraph("QA_MARKER_PARA_二")
    d.save(buf)
    return buf.getvalue()


def main() -> int:
    failed = 0

    def ok(name: str, cond: bool, detail: str = "") -> None:
        nonlocal failed
        if cond:
            print(f"[PASS] {name}")
        else:
            print(f"[FAIL] {name} {detail}")
            failed += 1

    with TestClient(app) as client:
        # 1) 首页
        r = client.get("/")
        ok("GET / 返回 200", r.status_code == 200 and "论文" in r.text)

        # 2) 静态资源
        r = client.get("/static/app.js")
        ok("GET /static/app.js", r.status_code == 200 and "upload-docx" in r.text)

        # 3) 诊断
        r = client.get("/api/diagnostics")
        ok("GET /api/diagnostics 200", r.status_code == 200)
        if r.status_code == 200:
            j = r.json()
            ok("diagnostics 含 thesis_exists", "thesis_exists" in j)

        # 4) 备份论文.txt
        backup: bytes | None = None
        if THESIS_PATH.exists():
            backup = THESIS_PATH.read_bytes()

        try:
            # 5) 上传 docx
            docx_bytes = make_minimal_docx_bytes()
            files = {
                "file": (
                    "qa_smoke.docx",
                    docx_bytes,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            }
            r = client.post("/api/upload-docx", files=files)
            ok("POST /api/upload-docx 成功", r.status_code == 200)
            if r.status_code == 200:
                body = r.json()
                ok("upload 返回 ok", body.get("ok") is True)
                ok("upload 返回 chars", isinstance(body.get("chars"), int) and body["chars"] > 0)

            text = THESIS_PATH.read_text(encoding="utf-8")
            ok("论文.txt 含 QA 标记", "QA_MARKER_PARA_一" in text and "QA_MARKER_PARA_二" in text)

            # 6) 错误：非 docx
            r = client.post(
                "/api/upload-docx",
                files={"file": ("x.txt", b"hello", "text/plain")},
            )
            ok("错误扩展名 400", r.status_code == 400)

            # 7) 错误：伪 zip
            r = client.post(
                "/api/upload-docx",
                files={"file": ("fake.docx", b"PK\x03\x04not a real docx", "application/octet-stream")},
            )
            ok("无效 docx 400", r.status_code == 400)

            # 8) 空内容 docx（仅空白段落）
            from docx import Document

            buf = io.BytesIO()
            d = Document()
            d.add_paragraph("   ")
            d.save(buf)
            r = client.post(
                "/api/upload-docx",
                files={"file": ("emptyish.docx", buf.getvalue(), "application/octet-stream")},
            )
            ok("无有效文本 400", r.status_code == 400)

        finally:
            if backup is not None:
                THESIS_PATH.write_bytes(backup)
                print("[INFO] 已恢复 论文.txt 备份")
            else:
                if THESIS_PATH.exists():
                    THESIS_PATH.unlink()
                    print("[INFO] 已删除测试生成的 论文.txt（原不存在）")

    print("[INFO] 改写流程 /api/start 需配置 API Key 后手动或集成测试 tests/integration_llm_env.py。")

    print("---")
    if failed:
        print(f"结果: {failed} 项失败")
        return 1
    print("结果: 全部通过")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
